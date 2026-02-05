"""
DOCX document processor for Word files.
"""

import io
import logging
from pathlib import Path
from typing import Generator
import uuid

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.shared import Inches
from docx.table import Table
from docx.text.paragraph import Paragraph

from .base import BaseFileProcessor, ContentChunk, ContentType
from settings import settings

logger = logging.getLogger("OfficeTranslator.DocxProcessor")


class DocxProcessor(BaseFileProcessor):
    """Processor for Microsoft Word (.docx) documents."""
    
    SUPPORTED_EXTENSIONS = ["docx"]
    
    def __init__(self):
        super().__init__()
        self._document: Document = None
        self._paragraph_map: dict[str, Paragraph] = {}
        self._image_map: dict[str, tuple] = {}  # id -> (run, inline_shape_index)
    
    def load(self, file_path: str | Path) -> None:
        """Load a DOCX file."""
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")
        
        if path.suffix.lower() != ".docx":
            raise ValueError(f"Not a DOCX file: {path}")
        
        try:
            self._document = Document(path)
            self._file_path = path
            self._paragraph_map.clear()
            self._image_map.clear()
            logger.info(f"Loaded DOCX: {path.name}")
        except PackageNotFoundError as e:
            raise ValueError(f"Invalid DOCX file: {e}")
    
    def extract_content_generator(self) -> Generator[ContentChunk, None, None]:
        """Extract content from DOCX paragraphs and tables."""
        self._validate_loaded()
        
        chunk_index = 0
        
        # Process paragraphs in body
        for para_idx, paragraph in enumerate(self._document.paragraphs):
            if paragraph.text.strip():
                chunk_id = f"para_{para_idx}"
                self._paragraph_map[chunk_id] = paragraph
                
                yield ContentChunk(
                    id=chunk_id,
                    content_type=ContentType.TEXT,
                    text=paragraph.text,
                    location=f"Paragraph {para_idx + 1}",
                    metadata={
                        "style": paragraph.style.name if paragraph.style else None,
                        "has_formatting": self._has_formatting(paragraph),
                    }
                )
                chunk_index += 1
            
            # Check for inline images
            for run in paragraph.runs:
                for inline_idx, inline in enumerate(run._element.findall('.//a:blip', 
                    {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})):
                    try:
                        image_data = self._extract_inline_image(inline)
                        if image_data:
                            img_id = f"img_{para_idx}_{inline_idx}"
                            self._image_map[img_id] = (run, inline_idx)
                            
                            yield ContentChunk(
                                id=img_id,
                                content_type=ContentType.IMAGE,
                                image_data=image_data,
                                location=f"Image in Paragraph {para_idx + 1}",
                            )
                            chunk_index += 1
                    except Exception as e:
                        logger.warning(f"Failed to extract image: {e}")
        
        # Process tables
        for table_idx, table in enumerate(self._document.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if cell_text:
                        chunk_id = f"table_{table_idx}_r{row_idx}_c{cell_idx}"
                        # Store reference to cell paragraphs
                        self._paragraph_map[chunk_id] = cell.paragraphs[0] if cell.paragraphs else None
                        
                        yield ContentChunk(
                            id=chunk_id,
                            content_type=ContentType.TABLE_CELL,
                            text=cell_text,
                            location=f"Table {table_idx + 1}, Row {row_idx + 1}, Cell {cell_idx + 1}",
                            metadata={
                                "table_index": table_idx,
                                "row": row_idx,
                                "col": cell_idx,
                            }
                        )
                        chunk_index += 1
        
        logger.info(f"Extracted {chunk_index} content chunks from DOCX")
    
    def apply_translation(
        self,
        chunk_id: str,
        translated_content: str | bytes,
    ) -> None:
        """Apply translation to a paragraph or table cell."""
        self._validate_loaded()
        
        if chunk_id.startswith("img_"):
            # Handle image replacement
            if chunk_id in self._image_map and isinstance(translated_content, bytes):
                # Note: Replacing inline images in docx is complex
                # For now, we log a warning - full implementation would need
                # to manipulate the document's relationship parts
                logger.warning(f"Image replacement in DOCX not yet implemented for {chunk_id}")
            return
        
        if chunk_id not in self._paragraph_map:
            logger.warning(f"Chunk not found: {chunk_id}")
            return
        
        paragraph = self._paragraph_map[chunk_id]
        if paragraph is None:
            return
        
        if settings.preserve_formatting and self._has_formatting(paragraph):
            self._apply_with_formatting(paragraph, str(translated_content))
        else:
            # Simple replacement - clears formatting but is reliable
            paragraph.clear()
            paragraph.add_run(str(translated_content))
        
        logger.debug(f"Applied translation to {chunk_id}")
    
    def save(self, output_path: str | Path) -> None:
        """Save the modified document."""
        self._validate_loaded()
        
        path = Path(output_path)
        self._document.save(path)
        logger.info(f"Saved DOCX: {path.name}")
    
    def _has_formatting(self, paragraph: Paragraph) -> bool:
        """Check if paragraph has inline formatting (bold/italic)."""
        for run in paragraph.runs:
            if run.bold or run.italic or run.underline:
                return True
        return False
    
    def _apply_with_formatting(self, paragraph: Paragraph, translated_text: str) -> None:
        """
        Apply translation while attempting to preserve formatting.
        
        This is a simplified implementation - full formatting preservation
        would require parsing pseudo-tags from the translation.
        """
        # For now, preserve the first run's formatting
        if paragraph.runs:
            first_run = paragraph.runs[0]
            bold = first_run.bold
            italic = first_run.italic
            underline = first_run.underline
            font_name = first_run.font.name
            font_size = first_run.font.size
            
            paragraph.clear()
            new_run = paragraph.add_run(translated_text)
            new_run.bold = bold
            new_run.italic = italic
            new_run.underline = underline
            if font_name:
                new_run.font.name = font_name
            if font_size:
                new_run.font.size = font_size
        else:
            paragraph.clear()
            paragraph.add_run(translated_text)
    
    def _extract_inline_image(self, blip_element) -> bytes | None:
        """Extract image bytes from an inline blip element."""
        try:
            rId = blip_element.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
            if rId:
                image_part = self._document.part.related_parts.get(rId)
                if image_part:
                    return image_part.blob
        except Exception as e:
            logger.debug(f"Image extraction failed: {e}")
        return None
